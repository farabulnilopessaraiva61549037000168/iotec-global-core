import sys

try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass
import os
from docx import Document

def gerar_documentacao(caminho, nome_doc):
    doc = Document()
    doc.add_heading('Blueprint do Sistema', 0)

    for raiz, dirs, arquivos in os.walk(caminho):
        nivel = raiz.replace(caminho, '').count(os.sep)
        indent = '    ' * nivel
        doc.add_paragraph(f'{indent}{os.path.basename(raiz)}/')
        sub_indent = '    ' * (nivel + 1)
        for arquivo in arquivos:
            doc.add_paragraph(f'{sub_indent}{arquivo}')

    doc.save(f'{nome_doc}.docx')
    print(f'DocumentaÃƒÆ'Ã†â€™ÃƒÆ'Ã†â€™o salva em {nome_doc}.docx')

# ÃƒÆ'Ã‚Â°Ãƒâ€¦Ã‚Â¸ÃƒÂ¢Ã¢â€šÂ¬Ã‚ÂÃƒâ€šÃ‚Â¥ EXECUTAR
gerar_documentacao('C:/MeuSistema', 'Blueprint_Sistema')



